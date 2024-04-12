from docx import Document
import google.generativeai as genai

api_key = "AIzaSyB-HHANZ8sEMS0q53peuzQr7U69ureKJRQ"
genai.configure(api_key=api_key)
model = genai.GenerativeModel('gemini-pro')


def translate(file_name, translate_language):
    doc = Document(file_name)

    original_language = ""
    for paragraph in doc.paragraphs:
        if len(paragraph.text) > 0:
            original_language += paragraph.text + "\n"

    original_language = original_language.split("\n")

    translated = []

    for sentence in original_language:
        translate = model.generate_content("Translate to english: " + sentence)
        translated.append(translate.text)

    return original_language, translated

def find_replace_text(org_doc, find_texts, replace_texts, new_doc):
    for find_text, replace_text in zip(find_texts, replace_texts):
        for paragraph in org_doc.paragraphs:
            if find_text == paragraph.text:
                if paragraph.runs:
                    run = paragraph.runs[0]
                    new_paragraph = new_doc.add_paragraph("")
                    new_paragraph.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
                    new_paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
                    new_paragraph = new_paragraph.add_run(replace_text)
                    new_paragraph.bold = run.bold
                    new_paragraph.font.size = run.font.size
                    new_paragraph.font.name = run.font.name
                    new_paragraph.font.color.rgb = run.font.color.rgb
                    new_paragraph.font.highlight_color = run.font.highlight_color
                    new_paragraph.font.shadow = run.font.shadow
                    new_paragraph.italic = run.italic
                    new_paragraph.underline = run.underline
                else:
                    continue


doc_name = input("Enter the file name: ")
language = input("Enter the language: ")
original_language, translated = translate(doc_name, translate_language=language)
doc = Document(doc_name)

exported_doc = Document()
find_replace_text(doc, original_language, translated, exported_doc)

exported_doc.save("exported.docx")
